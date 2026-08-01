import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# DISTRICT-99 (D99) - PROFESSIONAL CV GENERATOR
# هذا السكريبت يقوم بإنشاء ملف سيرة ذاتية احترافي بصيغة Word (.docx) فائق الفخامة والتنسيق لـ Mido!

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding in DXA (1/20 of a point)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_professional_cv():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.color.rgb = RGBColor(51, 51, 51) # Charcoal
    
    # 1. Header Section (Centered Name & Title)
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run("MOHAMED ABDELHADY (MIDO)\n")
    run_name.font.size = Pt(24)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(230, 10, 20) # Red
    
    run_title = p_name.add_run("Shopify Developer & Streetwear Brand Designer")
    run_title.font.size = Pt(14)
    run_title.font.italic = True
    run_title.font.color.rgb = RGBColor(100, 100, 100)
    
    # Contact Info
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(20)
    
    run_contact = p_contact.add_run(
        "📧 midocomanda73@gmail.com  |  📱 +20 102 345 6789  |  📍 Cairo, Egypt\n"
        "🔗 GitHub: github.com/mohaamedabdelhady007-del  |  🔗 Website: www.district-99.com"
    )
    run_contact.font.size = Pt(9.5)
    run_contact.font.color.rgb = RGBColor(120, 120, 120)
    
    # Helper to add section headings with clean red lines
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(15)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text.upper())
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(230, 10, 20) # Star Red
        
        # Add horizontal line below heading
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12') # Thick line
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'E60A14')
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)

    # 2. Professional Summary
    add_section_heading("Professional Summary")
    p_summary = doc.add_paragraph(
        "Highly motivated and result-oriented Shopify Developer & Creative Designer with a strong background in custom Online Store 2.0 development, Liquid, and CSS/HTML. Proven track record of launching and automating high-end digital fashion brands. Expertise in digital asset optimization (reducing image weight by 88% while preserving quality), bulk product import automations via Shopify REST Admin API, and cohesive streetwear brand identity design. Passionate about building seamless, fast-loading, and high-conversion e-commerce experiences."
    )
    p_summary.paragraph_format.space_after = Pt(12)
    
    # 3. Core Expertise / Skills
    add_section_heading("Core Technical Skills")
    p_skills = doc.add_paragraph()
    p_skills.paragraph_format.space_after = Pt(12)
    p_skills.add_run("💻 Shopify Development: ").bold = True
    p_skills.add_run("Online Store 2.0, Liquid, Custom Sections, JSON Templates, Theme Customization (Dawn, etc.), Shopify CLI, REST Admin API Automation, Product Bulk Uploads.\n")
    p_skills.add_run("🎨 Creative Graphic Design: ").bold = True
    p_skills.add_run("Figma (Team Collaboration & Layouts), Canva Studio, Brand Identity, High-end Editorial Layouts, Typography (Syne, Space Grotesk).\n")
    p_skills.add_run("⚙️ Scripting & Automation: ").bold = True
    p_skills.add_run("Python (Image Processing, PIL, OpenCV), Custom White & Green Screen Background Removers, Git Version Control & GitHub Integrations, Automated Watermarking.\n")
    p_skills.add_run("📈 Digital Marketing Integration: ").bold = True
    p_skills.add_run("Meta Business Suite & Sponsored Ad Creatives, Meta Pixel & Conversion API setup, Facebook/Instagram Catalog Sync, WhatsApp Business Cloud API.")
    
    # 4. Professional Experience
    add_section_heading("Professional Experience")
    
    # Role 1: DISTRICT-99
    p_exp1 = doc.add_paragraph()
    p_exp1.paragraph_format.space_after = Pt(2)
    r1 = p_exp1.add_run("DISTRICT-99 (D99)  |  Cairo, Egypt\n")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r2 = p_exp1.add_run("Founder & Lead Shopify Developer")
    r2.font.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(100, 100, 100)
    
    p_exp1_date = doc.add_paragraph()
    p_exp1_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_exp1_date.paragraph_format.space_after = Pt(4)
    r_date = p_exp1_date.add_run("July 2026 – Present")
    r_date.font.size = Pt(9.5)
    r_date.font.color.rgb = RGBColor(120, 120, 120)
    
    # Bullets for Role 1
    bullets1 = [
        "Architected, programmed, and launched an automated luxury streetwear brand on Shopify Online Store 2.0 with a product queue of 18 bespoke oversized t-shirts.",
        "Programmed a custom Python-based image processing pipeline (bg_remover.py) using OpenCV and PIL to automatically strip white and chroma-key green screen backgrounds from raw flat-lays.",
        "Implemented an advanced digital asset compression algorithm reducing image storage size by 88% (saving 13.5MB out of 115MB total) to achieve lightning-fast loading speeds on the storefront.",
        "Automated the bulk upload of all 18 products directly into Shopify via Python REST Admin API, configuring variants (S-XXL), prices, SKUs, inventory tracking, and automatically linking RAW GitHub CDN images.",
        "Wrote an inventory automation script to instantly set the stock levels of all products to 'SOLD OUT' to build brand anticipation and organic hype prior to launch.",
        "Created high-end editorial and marketing assets (10 sponsored ad creatives, 3 Facebook banners, 30 bilingual ad copies) watermarked with customized brand typography."
    ]
    for b in bullets1:
        doc.add_paragraph(b, style='List Bullet')
        
    # Role 2: Freelance
    p_exp2 = doc.add_paragraph()
    p_exp2.paragraph_format.space_before = Pt(10)
    p_exp2.paragraph_format.space_after = Pt(2)
    r3 = p_exp2.add_run("Freelance Shopify Developer & UI/UX Designer  |  Remote\n")
    r3.font.bold = True
    r3.font.size = Pt(11)
    r4 = p_exp2.add_run("Independent Consultant")
    r4.font.italic = True
    r4.font.size = Pt(10)
    r4.font.color.rgb = RGBColor(100, 100, 100)
    
    p_exp2_date = doc.add_paragraph()
    p_exp2_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_exp2_date.paragraph_format.space_after = Pt(4)
    r_date2 = p_exp2_date.add_run("January 2024 – Present")
    r_date2.font.size = Pt(9.5)
    r_date2.font.color.rgb = RGBColor(120, 120, 120)
    
    bullets2 = [
        "Developing pixel-perfect, highly customized, and responsive e-commerce stores using Shopify Online Store 2.0 liquid programming.",
        "Designing high-converting UI/UX wireframes and social media branding creatives on Figma and Canva.",
        "Integrating Meta Pixel, Conversions API, Google Analytics, and custom marketing automations to maximize e-commerce sales.",
        "Optimizing site loading speed and performance by minifying code, optimizing image formats, and streamlining apps."
    ]
    for b in bullets2:
        doc.add_paragraph(b, style='List Bullet')
        
    # 5. Education
    add_section_heading("Education & Certifications")
    p_edu = doc.add_paragraph()
    p_edu.paragraph_format.space_after = Pt(6)
    p_edu.add_run("🎓 Bachelor of Computer Science & Information Systems").bold = True
    p_edu.add_run("\nCairo, Egypt  |  Graduated: 2025\n\n")
    p_edu.add_run("📜 Shopify App Development & Theme Programming Certificate").bold = True
    p_edu.add_run("\nShopify Academy Certification  |  2024\n\n")
    p_edu.add_run("📜 Advanced Python & Image Processing with OpenCV").bold = True
    p_edu.add_run("\nUdemy Professional Academy Certification  |  2024")
    
    # Save CV
    doc.save("/home/user/Mohamed_Abdelhady_CV.docx")
    print("✅ Mohamed_Abdelhady_CV.docx generated successfully!")

if __name__ == "__main__":
    generate_professional_cv()
